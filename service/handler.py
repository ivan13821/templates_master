from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT










class Handler:
    """Осуществляет бизнес логику и обрабтку тела запрса"""






    @staticmethod
    def add_table(mass: dict, file_name:str, style: str = 'Table Grid', alignment: str = WD_TABLE_ALIGNMENT.LEFT, autofit: bool = True):


        doc = Document(f"templates/{file_name}.docx")

        table = doc.add_table(rows=0, cols=0)
        table.style = style
        table.autofit = autofit
        table.alignment = alignment


        for i in sorted(list(map(int, mass.keys()))):
            table.add_column(int(mass[str(i)]["size"]))
        

        #максимальное количество строк
        max_lenght = 0
        for i in mass.keys():
            #количество строк
            lenght = 0
            col = mass[i]["context"]
            for text in col:
                lenght += 1
                if lenght > max_lenght:
                    table.add_row()
                    max_lenght += 1
                table.cell(row_idx=lenght-1, col_idx=int(i)-1).text = text


        doc.save("test.docx")





mass = {
            "1":{
                "size":"200000",
                "context":["Коровы", "нюрка"]},
            "2":{
                "size":"400000",
                "context":["Собаки", "Шарик"]},
            "3":{
                "size":"800000",
                "context":["Попугаи", "Кеша"]
            },
            "4":{
                "size":"1600000",
                "context":["Ящерицы", "Амбал"]
            },
            "5":{
                "size":"3200000",
                "context":["Кролики", "Валера"]
            },
            "6":{
                "size":"3200000",
                "context":["Кролики", "Валера"]
            },
            "7":{
                "size":"3200000",
                "context":["Кролики", "Валера"]
            },
            "8":{
                "size":"3200000",
                "context":["Кролики", "Валера"]
            }
        }





Handler.add_table(mass, file_name="Приказ_о_заявлении_для_налоговой", autofit=True)